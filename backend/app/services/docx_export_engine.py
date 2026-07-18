"""Word Export Engine — headings, paragraphs, tables, images, captions,
code boxes and references, built with python-docx from the article JSON."""
from __future__ import annotations

import io
import re

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x22, 0x28, 0x3F)
ACCENT = RGBColor(0x4F, 0x6B, 0xF0)
SOFT = RGBColor(0x6E, 0x77, 0x96)

_INLINE = re.compile(
    r"(\*\*.+?\*\*|\*[^*\n]+\*|`[^`\n]+`|\$[^$\n]+\$|⟦[^⟧]+⟧)")
_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
_SEPARATOR = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
_HEADING = re.compile(r"^(#{2,4})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def _add_runs(paragraph, text: str, base_size=11, mono=False):
    for chunk in _INLINE.split(text):
        if not chunk:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(base_size)
        if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
            run.text = chunk[2:-2]
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            run.text = chunk[1:-1]
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run.text = chunk[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 0.5)
        elif chunk.startswith("⟦") and chunk.endswith("⟧"):
            run.text = "[" + chunk[1:-1] + "]"
            run.font.superscript = True
            run.font.size = Pt(max(7, base_size - 3))
            run.font.color.rgb = ACCENT
        elif chunk.startswith("$") and chunk.endswith("$"):
            run.text = chunk[1:-1]
            run.italic = True
            run.font.name = "Cambria Math"
        else:
            run.text = chunk
            if mono:
                run.font.name = "Consolas"


def _add_code(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK


def _add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            text = row[ci] if ci < len(row) else ""
            _add_runs(para, text, base_size=10)
            if ri == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def _fetch_image(url: str) -> bytes | None:
    try:
        if url.startswith("data:"):
            import base64
            head, b64 = url.split(",", 1)
            raw = base64.b64decode(b64)
            if "svg" in head:
                try:
                    import cairosvg  # optional
                    return cairosvg.svg2png(bytestring=raw, output_width=1100)
                except Exception:
                    return None
            return raw
        with httpx.Client(timeout=20, follow_redirects=True) as c:
            r = c.get(url, headers={"User-Agent": "SEARCH-AI/1.0"})
            r.raise_for_status()
            return r.content
    except Exception:
        return None


def _add_image(doc: Document, img: dict):
    raw = _fetch_image(img.get("url", ""))
    if raw:
        try:
            doc.add_picture(io.BytesIO(raw), width=Inches(5.9))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            raw = None
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = "Figure" if raw else "Figure (see live article / PDF export)"
    run = cap.add_run(f"{label} {img.get('slot','')}: {img.get('caption','')}"
                      f"  ·  {img.get('source_label','')}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = SOFT


def _render_markdown(doc: Document, md: str):
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code(doc, "\n".join(code_lines))
            i += 1
            continue
        if _TABLE_ROW.match(line):
            rows = []
            while i < len(lines) and _TABLE_ROW.match(lines[i]):
                if not _SEPARATOR.match(lines[i]):
                    cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                    rows.append(cells)
                i += 1
            _add_table(doc, rows)
            continue
        h = _HEADING.match(line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(re.sub(r"[*`#]", "", h.group(2)).strip(), level=level)
            i += 1
            continue
        b = _BULLET.match(line)
        if b:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, b.group(1))
            i += 1
            continue
        nmatch = _NUMBERED.match(line)
        if nmatch:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, nmatch.group(1))
            i += 1
            continue
        if line.strip().startswith("$$"):
            math_lines = [line.strip().strip("$")]
            if not line.strip().endswith("$$") or line.strip() == "$$":
                i += 1
                while i < len(lines) and "$$" not in lines[i]:
                    math_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    math_lines.append(lines[i].replace("$$", ""))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" ".join(l.strip() for l in math_lines if l.strip()))
            run.italic = True
            run.font.name = "Cambria Math"
            i += 1
            continue
        if line.strip():
            p = doc.add_paragraph()
            _add_runs(p, line.strip())
        i += 1


def build_docx(article: dict) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(article.get("title", "SEARCH AI Article"), level=0)
    for run in title.runs:
        run.font.color.rgb = INK
    sub = doc.add_paragraph()
    run = sub.add_run(f"SEARCH AI · {article.get('layout','').replace('_',' ')} "
                      f"· topic: {article.get('topic','')}")
    run.font.size = Pt(9)
    run.font.color.rgb = SOFT

    doc.add_heading("Abstract", level=1)
    _render_markdown(doc, article.get("abstract", ""))
    doc.add_heading("Direct Answer", level=1)
    _render_markdown(doc, article.get("executive_answer", ""))
    takeaways = article.get("key_takeaways") or []
    if takeaways:
        doc.add_heading("Key Takeaways", level=1)
        for kt in takeaways:
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, str(kt))

    images = article.get("images", [])
    by_section: dict[str, list[dict]] = {}
    loose: list[dict] = []
    for img in images:
        (by_section.setdefault(img.get("section_id", ""), [])
         if img.get("section_id") else loose).append(img)

    sec_ids = [s.get("id") for s in article.get("sections", [])]
    for si, sec in enumerate(article.get("sections", [])):
        doc.add_heading(sec.get("title", ""), level=1)
        if sec.get("pull_quote"):
            pq = doc.add_paragraph()
            pq.paragraph_format.left_indent = Inches(0.35)
            run = pq.add_run(f"“{sec['pull_quote']}”")
            run.italic = True
            run.bold = True
            run.font.color.rgb = ACCENT
            run.font.size = Pt(12.5)
        _render_markdown(doc, sec.get("markdown", ""))
        for img in by_section.get(sec.get("id", ""), []):
            _add_image(doc, img)
        # distribute unassigned images evenly
        if loose and si == min(si, len(sec_ids) - 1) and si % 2 == 1:
            _add_image(doc, loose.pop(0))
    for img in loose:
        _add_image(doc, img)

    refs = article.get("references", [])
    if refs:
        doc.add_heading("References", level=1)
        for r in refs:
            p = doc.add_paragraph(style="List Number")
            bits = [r.get("title", "")]
            if r.get("year"):
                bits.append(f"({r['year']})")
            if r.get("doi"):
                bits.append(f"doi:{r['doi']}")
            if r.get("url"):
                bits.append(r["url"])
            run = p.add_run("  ".join(str(b) for b in bits if b))
            run.font.size = Pt(9.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
